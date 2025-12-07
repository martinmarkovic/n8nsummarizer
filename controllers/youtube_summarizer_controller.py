"""
YouTube Summarizer Controller v3.1 - Coordinates YouTube tab + models

Responsibilities:
    - Listen to YouTubeSummarizerTab UI events
    - Validate YouTube URLs
    - Call TranscribeModel for YouTube transcription
    - Call N8NModel for summarization
    - Handle file operations (export txt, docx, clipboard)
    - Forward transcript to Transcriber tab
    - Manage threading for blocking operations
    - Error handling and user feedback

New in v3.1:
    - "Summarize" button instead of "Transcribe"
    - Export .docx instead of .srt
    - Forward transcript to Transcriber tab output folder
    - Show notification when transcription is available

Controller is THIN - just coordinates, doesn't contain business logic.

Version: 3.1
Created: 2025-12-07 (v3.0)
Updated: 2025-12-07 (v3.1 - Summarize button, .docx export, Transcriber tab integration)
"""
import os
import threading
from datetime import datetime
from tkinter import filedialog
from models.transcribe_model import TranscribeModel
from models.n8n_model import N8NModel
from utils.logger import logger
from config import EXPORT_DIR

# Try to import pyperclip, fallback to tkinter if not available
try:
    import pyperclip
    PYPERCLIP_AVAILABLE = True
except ImportError:
    PYPERCLIP_AVAILABLE = False
    logger.warning("pyperclip not installed. Clipboard functionality will use tkinter fallback.")


class YouTubeSummarizerController:
    """
    Coordinates YouTube Summarizer tab UI and models.
    
    Flow:
    1. User enters YouTube URL
    2. Controller validates URL
    3. Calls TranscribeModel to transcribe YouTube video
    4. Saves transcript to disk (Transcriber output folder)
    5. Sends transcript to N8NModel for summarization
    6. N8NModel calls n8n webhook
    7. Displays summary in UI
    8. Forwards transcript to Transcriber tab
    9. User can export summary or copy
    """
    
    def __init__(self, view, transcriber_tab=None):
        """
        Initialize controller with view reference.
        
        Args:
            view: YouTubeSummarizerTab instance
            transcriber_tab: TranscriberTab instance (for forwarding transcript)
        """
        self.view = view
        self.transcriber_tab = transcriber_tab  # Reference to Transcriber tab (set by main)
        self.transcribe_model = TranscribeModel()
        self.n8n_model = N8NModel()
        
        # Wire up view callbacks
        self.view.on_summarize_clicked = self.handle_summarize_clicked
        self.view.on_export_txt_clicked = self.handle_export_txt
        self.view.on_export_docx_clicked = self.handle_export_docx
        self.view.on_copy_clipboard_clicked = self.handle_copy_clipboard
        
        # State tracking
        self.current_summary = None
        self.current_transcript = None
        self.current_youtube_title = None
        self.current_transcript_format = None
        
        logger.info("YouTubeSummarizerController initialized (v3.1)")
    
    def handle_summarize_clicked(self):
        """
        Handle Summarize button - starts background thread (renamed in v3.1).
        """
        logger.info("Summarize button clicked")
        
        # Get and validate URL
        youtube_url = self.view.get_youtube_url()
        if not youtube_url or youtube_url == "https://":
            self.view.show_error("Please enter a YouTube URL")
            return
        
        # Validate YouTube URL format
        if not TranscribeModel.validate_youtube_url(youtube_url):
            self.view.show_error(
                "Invalid YouTube URL.\n\n"
                "Supported formats:\n"
                "- https://youtube.com/watch?v=...\n"
                "- https://youtu.be/...\n"
                "- https://youtube.com/embed/..."
            )
            return
        
        # Get selected format
        format_ext = self.view.get_transcription_format()
        keep_formats = [format_ext]
        self.current_transcript_format = format_ext
        
        # Update UI
        self.view.set_input_status("Transcribing YouTube video...", "blue")
        self.view.set_summarize_button_enabled(False)
        self.view.set_export_buttons_enabled(False)
        self.view.set_summary_content("Transcribing YouTube video...\nThis may take a few minutes depending on video length.")
        
        # Start background thread
        thread = threading.Thread(
            target=self._transcribe_youtube_thread,
            args=(youtube_url, keep_formats),
            daemon=True
        )
        thread.start()
    
    def _transcribe_youtube_thread(self, youtube_url: str, keep_formats: list):
        """
        Background thread for YouTube transcription.
        
        Args:
            youtube_url: YouTube video URL
            keep_formats: File formats to keep
        """
        try:
            logger.info(f"Starting YouTube transcription: {youtube_url}")
            
            # Transcribe YouTube video
            success, transcript_content, error, metadata = self.transcribe_model.transcribe_youtube(
                youtube_url=youtube_url,
                device="cuda",
                keep_formats=keep_formats
            )
            
            if not success:
                error_msg = error or "Unknown error during transcription"
                self.view.root.after(0, self._on_transcribe_error, error_msg)
                return
            
            if not transcript_content:
                self.view.root.after(
                    0,
                    self._on_transcribe_error,
                    "No transcript generated from YouTube video"
                )
                return
            
            logger.info(f"Transcription successful: {len(transcript_content)} characters")
            
            # Store transcript for later
            self.current_transcript = transcript_content
            self.current_youtube_title = metadata.get('base_name', 'youtube_video') if metadata else 'youtube_video'
            
            # Send to n8n for summarization
            self.view.root.after(
                0,
                self._on_transcribe_success,
                transcript_content,
                metadata
            )
        
        except Exception as e:
            error_msg = f"Unexpected error during transcription: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self.view.root.after(0, self._on_transcribe_error, error_msg)
    
    def _on_transcribe_success(self, transcript_content: str, metadata: dict):
        """
        Handle successful transcription - proceed to summarization.
        
        Args:
            transcript_content: Transcript text
            metadata: Transcription metadata
        """
        logger.info("Transcription succeeded, sending to n8n for summarization...")
        
        self.view.set_input_status("Sending to n8n for summarization...", "blue")
        self.view.set_summary_content(
            f"Transcript received ({len(transcript_content)} characters)\n\n"
            f"Sending to n8n for summarization...\n\n"
            f"Please wait..."
        )
        
        # Start background thread for n8n summarization
        thread = threading.Thread(
            target=self._summarize_with_n8n_thread,
            args=(transcript_content, metadata),
            daemon=True
        )
        thread.start()
    
    def _on_transcribe_error(self, error_msg: str):
        """
        Handle transcription error.
        
        Args:
            error_msg: Error message
        """
        logger.error(f"Transcription error: {error_msg}")
        self.view.show_error(f"Transcription failed:\n\n{error_msg}")
        self.view.set_input_status("Transcription failed", "red")
        self.view.set_summarize_button_enabled(True)
        self.view.set_summary_content(
            f"Error during transcription:\n\n{error_msg}\n\n"
            f"Please try again."
        )
    
    def _summarize_with_n8n_thread(self, transcript_content: str, metadata: dict):
        """
        Background thread for n8n summarization.
        
        Args:
            transcript_content: Transcript text to summarize
            metadata: Metadata about transcript
        """
        try:
            logger.info("Sending transcript to n8n webhook...")
            
            # Send to n8n
            success, response_data, error = self.n8n_model.send_content(
                file_name=self.current_youtube_title or 'youtube_video',
                content=transcript_content,
                metadata={
                    'source': 'youtube',
                    'transcript_length': len(transcript_content)
                }
            )
            
            if not success:
                error_msg = error or "Unknown error from n8n"
                self.view.root.after(0, self._on_n8n_error, error_msg)
                return
            
            # Extract summary from response
            summary = self.n8n_model.extract_summary(response_data)
            
            if not summary:
                self.view.root.after(
                    0,
                    self._on_n8n_error,
                    "No summary returned from n8n"
                )
                return
            
            logger.info(f"Summary received: {len(summary)} characters")
            self.current_summary = summary
            
            # Display summary
            self.view.root.after(0, self._on_n8n_success)
        
        except Exception as e:
            error_msg = f"Unexpected error during summarization: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self.view.root.after(0, self._on_n8n_error, error_msg)
    
    def _on_n8n_success(self):
        """
        Handle successful summarization.
        
        Displays summary and notifies user that transcript is available in Transcriber tab.
        """
        logger.info("Summarization succeeded")
        
        self.view.set_summary_content(self.current_summary)
        self.view.set_input_status("Summarization complete!", "green")
        self.view.set_output_status("Summary ready. You can now export or copy.", "green")
        self.view.set_export_buttons_enabled(True)
        self.view.set_summarize_button_enabled(True)
        
        # Forward transcript to Transcriber tab (if available)
        if self.transcriber_tab and self.current_transcript:
            try:
                self._forward_transcript_to_transcriber()
                
                # Show notification to user
                self.view.show_info(
                    "Summarization Complete!",
                    f"YouTube video successfully transcribed and summarized!\n\n"
                    f"✓ Summary is displayed above\n"
                    f"✓ Full transcript ({self.current_transcript_format}) forwarded to Transcriber tab\n\n"
                    f"Switch to the Transcriber tab to view the full transcript."
                )
                logger.info("Transcript forwarded to Transcriber tab")
            except Exception as e:
                logger.warning(f"Could not forward transcript to Transcriber tab: {str(e)}")
        else:
            self.view.show_success("YouTube video transcribed and summarized successfully!")
    
    def _on_n8n_error(self, error_msg: str):
        """
        Handle n8n summarization error.
        
        Args:
            error_msg: Error message
        """
        logger.error(f"n8n summarization error: {error_msg}")
        self.view.show_error(f"Summarization failed:\n\n{error_msg}")
        self.view.set_input_status("Summarization failed", "red")
        self.view.set_output_status(f"Error: {error_msg}", "red")
        self.view.set_summarize_button_enabled(True)
    
    def _forward_transcript_to_transcriber(self):
        """
        Forward the transcript to Transcriber tab output area.
        
        This allows users to see the full transcript in the Transcriber tab.
        """
        if not self.transcriber_tab:
            logger.warning("Transcriber tab not available for transcript forwarding")
            return
        
        try:
            # Check if Transcriber tab has a transcript_text widget
            if hasattr(self.transcriber_tab, 'transcript_text'):
                # Set the transcript content in Transcriber tab
                transcript_text_widget = self.transcriber_tab.transcript_text
                transcript_text_widget.config(state=tk.NORMAL)
                transcript_text_widget.delete("1.0", tk.END)
                transcript_text_widget.insert(tk.END, self.current_transcript)
                transcript_text_widget.config(state=tk.DISABLED)
                
                logger.info("Transcript forwarded to Transcriber tab")
            else:
                logger.warning("Transcriber tab does not have transcript_text widget")
        except Exception as e:
            logger.error(f"Error forwarding transcript to Transcriber tab: {str(e)}")
    
    def handle_export_txt(self):
        """
        Export summary as .txt file.
        """
        if not self.current_summary:
            self.view.show_error("No summary to export. Please summarize a YouTube video first.")
            return
        
        # Determine filename
        default_filename = f"{self.current_youtube_title or 'youtube_summary'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        # Show file dialog
        file_path = filedialog.asksaveasfilename(
            title="Export Summary as .txt",
            defaultextension=".txt",
            initialdir=EXPORT_DIR,
            initialfile=default_filename,
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.current_summary)
            
            logger.info(f"Summary exported to: {file_path}")
            self.view.show_success(f"Summary exported successfully to:\n{file_path}")
            self.view.set_output_status(f"Exported to {os.path.basename(file_path)}", "green")
        
        except Exception as e:
            error_msg = f"Failed to export file: {str(e)}"
            logger.error(error_msg)
            self.view.show_error(error_msg)
    
    def handle_export_docx(self):
        """
        Export summary as .docx file (new in v3.1).
        
        Uses python-docx library to create a professional Word document.
        """
        if not self.current_summary:
            self.view.show_error("No summary to export. Please summarize a YouTube video first.")
            return
        
        # Check if python-docx is available
        try:
            from docx import Document
        except ImportError:
            self.view.show_error(
                "python-docx library not installed!\n\n"
                "Install it with: pip install python-docx"
            )
            return
        
        # Determine filename
        default_filename = f"{self.current_youtube_title or 'youtube_summary'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Show file dialog
        file_path = filedialog.asksaveasfilename(
            title="Export Summary as .docx",
            defaultextension=".docx",
            initialdir=EXPORT_DIR,
            initialfile=default_filename,
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            doc.add_heading(f"YouTube Summary: {self.current_youtube_title}", 0)
            
            # Add metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Source: YouTube Video", style='Heading 3')
            
            # Add summary content
            doc.add_heading("Summary", 1)
            doc.add_paragraph(self.current_summary)
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"Summary exported as DOCX to: {file_path}")
            self.view.show_success(f"Summary exported successfully as .docx to:\n{file_path}")
            self.view.set_output_status(f"Exported to {os.path.basename(file_path)}", "green")
        
        except Exception as e:
            error_msg = f"Failed to export DOCX file: {str(e)}"
            logger.error(error_msg)
            self.view.show_error(error_msg)
    
    def handle_copy_clipboard(self):
        """
        Copy summary to clipboard.
        
        Uses pyperclip if available, otherwise uses tkinter clipboard.
        """
        if not self.current_summary:
            self.view.show_error("No summary to copy. Please summarize a YouTube video first.")
            return
        
        try:
            if PYPERCLIP_AVAILABLE:
                # Use pyperclip
                pyperclip.copy(self.current_summary)
                logger.info("Summary copied to clipboard (pyperclip)")
            else:
                # Fallback to tkinter clipboard
                self.view.root.clipboard_clear()
                self.view.root.clipboard_append(self.current_summary)
                self.view.root.update()
                logger.info("Summary copied to clipboard (tkinter fallback)")
            
            self.view.show_success(f"Summary copied to clipboard ({len(self.current_summary)} characters)")
            self.view.set_output_status("Copied to clipboard", "green")
        
        except Exception as e:
            error_msg = f"Failed to copy to clipboard: {str(e)}"
            logger.error(error_msg)
            self.view.show_error(error_msg)


# Import tk for transcript forwarding
import tkinter as tk
