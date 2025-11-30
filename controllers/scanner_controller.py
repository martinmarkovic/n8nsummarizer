"""
Scanner Controller - with webhook override, threading, exports, theme toggle, and auto-export
"""
import os
import threading
from datetime import datetime
from tkinter import filedialog, messagebox
from models.file_scanner import FileScanner
from models.http_client import HTTPClient
from utils.logger import logger
from config import EXPORT_DIR


class ScannerController:
    """Controller with export, theme, and auto-export functionality"""
    
    def __init__(self, view):
        self.view = view
        self.file_scanner = FileScanner()
        self.http_client = HTTPClient()
        
        # Setup view callbacks
        self.view.on_file_selected = self.handle_file_selected
        self.view.on_send_clicked = self.handle_send_clicked
        self.view.on_clear_clicked = self.handle_clear_clicked
        self.view.on_export_txt = self.handle_export_txt
        self.view.on_export_docx = self.handle_export_docx
        self.view.on_theme_toggle = self.handle_theme_toggle
        
        logger.info("ScannerController initialized with export, theme, and auto-export support")
    
    def handle_file_selected(self, file_path):
        """Handle file selection from view"""
        logger.info(f"File selected: {file_path}")
        
        success, content, error = self.file_scanner.read_file(file_path)
        
        if success:
            self.view.set_file_path(file_path)
            self.view.set_content(content)
            
            file_info = self.file_scanner.get_file_info()
            self.view.set_file_info(file_info)
            
            self.view.set_status(f"File loaded: {file_info['name']}")
        else:
            self.view.show_error(f"Failed to load file: {error}")
            self.view.set_file_path(None)
            self.view.set_content("")
            self.view.set_file_info(None)
    
    def handle_send_clicked(self):
        """Handle send button click - starts background thread for n8n request"""
        logger.info("Send button clicked")
        
        if not self.file_scanner.current_file:
            self.view.show_error("No file loaded. Please select a file first.")
            return
        
        content = self.view.get_content()
        
        if not content or not content.strip():
            self.view.show_error("File content is empty. Cannot send empty content.")
            return
        
        webhook_override = self.view.get_webhook_override()
        gui_webhook_url = webhook_override['custom_url']
        
        if not gui_webhook_url or not gui_webhook_url.strip():
            self.view.show_error("Webhook URL in GUI is empty! Please enter a valid webhook URL.")
            return
        
        gui_webhook_url = gui_webhook_url.strip()
        
        if webhook_override['override']:
            try:
                self._save_webhook_to_env(gui_webhook_url)
                logger.info(f"Saved webhook to .env: {gui_webhook_url}")
            except Exception as e:
                logger.error(f"Failed to save webhook to .env: {e}")
                self.view.show_error(f"Warning: Could not save to .env file: {e}")
        
        file_info = self.file_scanner.get_file_info()
        status_msg = f"Sending request to n8n...\n\nWebhook: {gui_webhook_url}\nFile: {file_info['name']}\nSize: {file_info['size_kb']:.2f} KB\n\nWaiting for response..."
        self.view.display_response(status_msg)
        
        self.view.set_status("Sending to n8n and waiting for response...")
        self.view.show_loading(True)
        
        thread = threading.Thread(
            target=self._send_to_n8n_thread,
            args=(file_info, content, gui_webhook_url, webhook_override['override']),
            daemon=True
        )
        thread.start()
    
    def _send_to_n8n_thread(self, file_info, content, webhook_url, saved_to_env):
        """Background thread function to send request to n8n"""
        try:
            self.http_client.webhook_url = webhook_url
            logger.info(f"Using webhook from GUI: {webhook_url}")
            
            success, response_data, error = self.http_client.send_to_n8n(
                file_name=file_info['name'],
                content=content,
                metadata={
                    'size_bytes': file_info['size'],
                    'lines': file_info['lines']
                }
            )
            
            if success:
                summary = self._extract_summary(response_data)
                saved_msg = " (saved to .env)" if saved_to_env else ""
                
                if summary:
                    self.view.root.after(0, self._on_success_with_summary, summary, file_info['name'], saved_msg)
                else:
                    self.view.root.after(0, self._on_success_no_summary, file_info['name'], saved_msg)
            else:
                self.view.root.after(0, self._on_error, error)
        
        except Exception as e:
            error_msg = f"Unexpected error: {str(e)}"
            logger.error(error_msg)
            self.view.root.after(0, self._on_error, error_msg)
    
    def _on_success_with_summary(self, summary, file_name, saved_msg):
        """Handle successful summarization response"""
        self.view.display_response(summary)
        
        # Check if auto-export is enabled
        export_prefs = self.view.get_export_preferences()
        if export_prefs['auto_export']:
            logger.info("Auto-export enabled - exporting both .txt and .docx")
            self._auto_export_response(summary, silent=True)
            success_msg = f"Summarization received and auto-exported for '{file_name}'!{saved_msg}"
        else:
            success_msg = f"Summarization received for '{file_name}'!{saved_msg}"
        
        self.view.show_success(success_msg)
        logger.info("Summarization successfully received")
        self.view.set_status("Ready")
        self.view.show_loading(False)
    
    def _auto_export_response(self, response_content, silent=False):
        """Auto-export response as both .txt and .docx"""
        export_prefs = self.view.get_export_preferences()
        
        # Determine base filename - use original filename + "Summary" or timestamp
        if export_prefs['original_basename']:
            base_filename = f"{export_prefs['original_basename']}_Summary"
            logger.info(f"Using smart filename: {base_filename}")
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"n8n_response_{timestamp}"
            logger.info(f"Using timestamp filename: {base_filename}")
        
        # Determine save directory
        if export_prefs['use_original_location'] and export_prefs['original_directory']:
            save_dir = export_prefs['original_directory']
            logger.info(f"Auto-export to original location: {save_dir}")
        else:
            save_dir = EXPORT_DIR
            logger.info(f"Auto-export to default location: {save_dir}")
        
        # Export .txt
        txt_path = os.path.join(save_dir, f"{base_filename}.txt")
        self._save_txt_file(txt_path, response_content, silent=True)
        
        # Export .docx
        docx_path = os.path.join(save_dir, f"{base_filename}.docx")
        self._save_docx_file(docx_path, response_content, silent=True)
        
        logger.info(f"Auto-exported to:\n  .txt: {txt_path}\n  .docx: {docx_path}")
    
    def _on_success_no_summary(self, file_name, saved_msg):
        self.view.display_response("Request sent successfully.\n\nNo summary returned from n8n workflow.")
        self.view.show_success(f"Successfully sent '{file_name}' to n8n!{saved_msg}")
        logger.info("File successfully sent to n8n (no summary in response)")
        self.view.set_status("Ready")
        self.view.show_loading(False)
    
    def _on_error(self, error_msg):
        self.view.display_response(f"Error occurred:\n\n{error_msg}")
        self.view.show_error(f"Failed to send to n8n:\n{error_msg}")
        logger.error(f"Send failed: {error_msg}")
        self.view.set_status("Ready")
        self.view.show_loading(False)
    
    def handle_export_txt(self, manual_call=True):
        """Export response content as .txt file (manual export)"""
        response_content = self.view.get_response_content()
        
        if not response_content or not response_content.strip():
            self.view.show_error("No response content to export!")
            return
        
        # Get export preferences
        export_prefs = self.view.get_export_preferences()
        
        # Determine filename - use original filename + "Summary" or timestamp
        if export_prefs['original_basename']:
            default_filename = f"{export_prefs['original_basename']}_Summary.txt"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"n8n_response_{timestamp}.txt"
        
        # Determine initial directory
        if export_prefs['use_original_location'] and export_prefs['original_directory']:
            initial_dir = export_prefs['original_directory']
            logger.info(f"Using original file location for export: {initial_dir}")
            # Auto-save without dialog
            file_path = os.path.join(initial_dir, default_filename)
            self._save_txt_file(file_path, response_content)
        else:
            # Show file dialog
            file_path = filedialog.asksaveasfilename(
                title="Export Response as .txt",
                defaultextension=".txt",
                initialdir=EXPORT_DIR,
                initialfile=default_filename,
                filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
            )
            
            if file_path:
                self._save_txt_file(file_path, response_content)
    
    def _save_txt_file(self, file_path, content, silent=False):
        """Save content to .txt file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Exported response to: {file_path}")
            if not silent:
                self.view.show_success(f"Response exported successfully to:\n{file_path}")
                self.view.set_status(f"Exported to {os.path.basename(file_path)}")
        except Exception as e:
            error_msg = f"Failed to export .txt file: {str(e)}"
            logger.error(error_msg)
            if not silent:
                self.view.show_error(error_msg)
    
    def handle_export_docx(self):
        """Export response content as .docx file (manual export)"""
        response_content = self.view.get_response_content()
        
        if not response_content or not response_content.strip():
            self.view.show_error("No response content to export!")
            return
        
        try:
            from docx import Document
        except ImportError:
            self.view.show_error(
                "python-docx library not installed!\n\n"
                "Install it with: pip install python-docx"
            )
            return
        
        # Get export preferences
        export_prefs = self.view.get_export_preferences()
        
        # Determine filename - use original filename + "Summary" or timestamp
        if export_prefs['original_basename']:
            default_filename = f"{export_prefs['original_basename']}_Summary.docx"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"n8n_response_{timestamp}.docx"
        
        # Determine initial directory
        if export_prefs['use_original_location'] and export_prefs['original_directory']:
            initial_dir = export_prefs['original_directory']
            logger.info(f"Using original file location for export: {initial_dir}")
            # Auto-save without dialog
            file_path = os.path.join(initial_dir, default_filename)
            self._save_docx_file(file_path, response_content)
        else:
            # Show file dialog
            file_path = filedialog.asksaveasfilename(
                title="Export Response as .docx",
                defaultextension=".docx",
                initialdir=EXPORT_DIR,
                initialfile=default_filename,
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
            )
            
            if file_path:
                self._save_docx_file(file_path, response_content)
    
    def _save_docx_file(self, file_path, content, silent=False):
        """Save content to .docx file"""
        try:
            from docx import Document
            
            document = Document()
            document.add_heading('n8n Response Summary', 0)
            
            document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            document.add_paragraph()
            
            for line in content.split('\n'):
                document.add_paragraph(line)
            
            document.save(file_path)
            
            logger.info(f"Exported response to: {file_path}")
            if not silent:
                self.view.show_success(f"Response exported successfully to:\n{file_path}")
                self.view.set_status(f"Exported to {os.path.basename(file_path)}")
        except Exception as e:
            error_msg = f"Failed to export .docx file: {str(e)}"
            logger.error(error_msg)
            if not silent:
                self.view.show_error(error_msg)
    
    def handle_theme_toggle(self, theme):
        """Handle theme toggle - save preference to .env"""
        logger.info(f"Theme changed to: {theme}")
        
        try:
            self._save_theme_to_env(theme)
            logger.info(f"Saved theme preference to .env: {theme}")
        except Exception as e:
            logger.error(f"Failed to save theme to .env: {e}")
    
    def _save_theme_to_env(self, theme):
        env_file = '.env'
        env_lines = []
        theme_found = False
        
        if os.path.exists(env_file):
            with open(env_file, 'r', encoding='utf-8') as f:
                env_lines = f.readlines()
        
        new_lines = []
        for line in env_lines:
            if line.strip().startswith('APP_THEME='):
                new_lines.append(f'APP_THEME={theme}\n')
                theme_found = True
            else:
                new_lines.append(line)
        
        if not theme_found:
            new_lines.append(f'APP_THEME={theme}\n')
        
        with open(env_file, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
    
    def _save_webhook_to_env(self, webhook_url):
        env_file = '.env'
        env_lines = []
        webhook_found = False
        
        if os.path.exists(env_file):
            with open(env_file, 'r', encoding='utf-8') as f:
                env_lines = f.readlines()
        
        new_lines = []
        for line in env_lines:
            if line.strip().startswith('N8N_WEBHOOK_URL='):
                new_lines.append(f'N8N_WEBHOOK_URL={webhook_url}\n')
                webhook_found = True
            else:
                new_lines.append(line)
        
        if not webhook_found:
            new_lines.append(f'N8N_WEBHOOK_URL={webhook_url}\n')
        
        with open(env_file, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
    
    def _extract_summary(self, response_data):
        if response_data is None:
            return None
        
        if isinstance(response_data, str):
            return response_data
        
        if isinstance(response_data, dict):
            for key in ['summary', 'summarization', 'result', 'output', 'text', 'content']:
                if key in response_data:
                    return response_data[key]
            
            import json
            return json.dumps(response_data, indent=2)
        
        return str(response_data)
    
    def handle_clear_clicked(self):
        logger.info("Clear button clicked")
        self.file_scanner.clear()
        self.view.clear_all()
        self.view.set_status("Ready")
    
    def test_n8n_connection(self):
        logger.info("Testing n8n connection...")
        is_connected = self.http_client.test_connection()
        
        if is_connected:
            self.view.set_status("n8n server is reachable")
            logger.info("n8n connection test passed")
        else:
            self.view.show_error("Cannot reach n8n server. Check if it's running at the configured URL.")
            logger.warning("n8n connection test failed")
        
        return is_connected
