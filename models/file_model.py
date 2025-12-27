"""
File Model - Business logic for file operations

Responsibilities:
    - Read files (text, csv, json, .docx, .srt)
    - Extract file information (size, line count, character count)
    - Validate file format
    - Export to text/docx formats
    - Auto-detect file encoding (UTF-8, UTF-16, etc.)

New in v4.4.4:
    - Automatic encoding detection for .srt and text files
    - Handles UTF-16, UTF-8, Latin-1, and other encodings
    - Fixes issue where 181KB files showed as 78.7KB
    
This is PURE business logic - NO UI dependencies.
Can be tested independently without GUI.
"""
import os
import json
from datetime import datetime
from pathlib import Path
from config import DEFAULT_ENCODING, EXPORT_DIR
from utils.logger import logger
from utils.validators import validate_file, validate_content_not_empty


class FileModel:
    """Handles all file-related operations"""
    
    def read_file(self, file_path: str) -> tuple[bool, str, str]:
        """
        Read file content with support for multiple formats and automatic encoding detection.
        Supports: .txt, .srt, .json, .docx, and text-based formats.
        
        v4.4.4: Auto-detect encoding (UTF-8, UTF-16, UTF-16-LE, UTF-16-BE, etc.)
        This fixes files that show half their actual size (UTF-16 encoded).
        
        Args:
            file_path (str): Path to file to read
            
        Returns:
            tuple: (success: bool, content: str, error_msg: str or None)
            
        Example:
            >>> model = FileModel()
            >>> success, content, error = model.read_file('test.txt')
            >>> if success:
            ...     print(f"Read {len(content)} characters")
        """
        # Validate file exists and is readable
        valid, error = validate_file(file_path)
        if not valid:
            logger.error(f"File validation failed: {error}")
            return False, "", error
        
        try:
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # Handle different file types
            if ext == '.docx':
                content = self._read_docx(file_path)
            elif ext == '.json':
                content = self._read_json(file_path)
            elif ext == '.srt':
                content = self._read_srt(file_path)
            else:
                # Default: read as text with encoding detection
                content = self._read_text_with_encoding_detection(file_path)
            
            # Validate content is not empty
            valid, error = validate_content_not_empty(content)
            if not valid:
                logger.warning(f"File is empty: {file_path}")
                return False, "", error
            
            logger.info(f"Successfully read file: {file_path} ({ext})")
            return True, content, None
            
        except UnicodeDecodeError as e:
            error = f"File encoding error - try .txt format: {str(e)}"
            logger.error(error)
            return False, "", error
        except Exception as e:
            error = f"Error reading file: {str(e)}"
            logger.error(error)
            return False, "", error
    
    def _read_text_with_encoding_detection(self, file_path: str) -> str:
        """
        v4.4.4: Read text file with automatic encoding detection.
        
        Tries multiple encodings in order:
        1. UTF-8 (most common)
        2. UTF-8 with BOM
        3. UTF-16 (Windows, half the reported size)
        4. UTF-16-LE (Little Endian)
        5. UTF-16-BE (Big Endian)
        6. CP1252 (Windows Latin)
        7. Latin-1 (fallback, accepts any byte sequence)
        
        This fixes the issue where 181 KB files show as 78.7 KB!
        The 181 KB file is UTF-16 encoded:
          181,000 bytes ÷ 2 bytes per char = ~90,500 characters
          90,500 chars ÷ 1024 = 78.7 KB (what was shown before)
        
        Args:
            file_path (str): Path to file
            
        Returns:
            str: File content
        """
        # List of encodings to try, in order of likelihood
        encodings_to_try = [
            'utf-8',
            'utf-8-sig',  # UTF-8 with BOM
            'utf-16',
            'utf-16-le',
            'utf-16-be',
            'cp1252',  # Windows Latin
            'iso-8859-1',  # Latin-1, accepts any byte sequence
        ]
        
        file_size = os.path.getsize(file_path)
        logger.debug(f"Reading {file_path} - File size: {file_size} bytes")
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                # Success! Log which encoding worked
                content_size = len(content)
                logger.debug(f"Successfully read with encoding '{encoding}': {content_size} characters")
                logger.debug(f"Size ratio: {file_size} bytes → {content_size} chars (ratio: {file_size/max(1,content_size):.2f})")
                
                # If UTF-16, check if we got roughly half the file size
                if encoding.startswith('utf-16') and file_size > content_size * 1.8:
                    logger.info(f"⚠️ Encoding '{encoding}': File was {file_size} bytes, content {content_size} chars")
                    logger.info(f"   This is expected for UTF-16 (uses 2+ bytes per character)")
                
                return content
            except (UnicodeDecodeError, UnicodeError, LookupError) as e:
                logger.debug(f"Encoding '{encoding}' failed: {str(e)[:100]}")
                continue
            except Exception as e:
                logger.debug(f"Unexpected error with '{encoding}': {str(e)[:100]}")
                continue
        
        # Fallback: If all else fails, read as Latin-1 (it accepts any byte sequence)
        logger.warning(f"Could not decode with standard encodings, using Latin-1 fallback")
        with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
            return f.read()
    
    def _read_docx(self, file_path: str) -> str:
        """Read .docx file and extract text"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            logger.warning("python-docx not installed")
            raise ImportError("python-docx library required to read DOCX files. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    
    def _read_json(self, file_path: str) -> str:
        """Read .json file and pretty-print it"""
        with open(file_path, 'r', encoding=DEFAULT_ENCODING) as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=False)
    
    def _read_srt(self, file_path: str) -> str:
        """Read .srt subtitle file and extract text only (skip timecodes)"""
        # v4.4.4: Use encoding detection for .srt files too
        content = self._read_text_with_encoding_detection(file_path)
        
        # Parse SRT: extract only subtitle text (skip numbers and timecodes)
        lines = content.split('\n')
        subtitles = []
        for line in lines:
            # Skip empty lines, numbers, and timecode lines
            line = line.strip()
            if line and not line.isdigit() and '-->' not in line:
                subtitles.append(line)
        
        return '\n'.join(subtitles)
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Extract file metadata.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: File information
                {
                    'name': filename,
                    'path': full_path,
                    'size_kb': size in KB,
                    'lines': line count,
                    'characters': total characters,
                    'characters_no_spaces': without whitespace
                }
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return {}
            
            # Read file to count lines/characters
            success, content, error = self.read_file(file_path)
            if not success:
                return {}
            
            file_stat = os.stat(file_path)
            
            return {
                'name': os.path.basename(file_path),
                'path': file_path,
                'size': file_stat.st_size,
                'size_kb': file_stat.st_size / 1024,
                'lines': len(content.split('\n')),
                'characters': len(content),
                'characters_no_spaces': len(content.replace(' ', '').replace('\n', ''))
            }
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {}
    
    def export_txt(self, content: str, filepath: str) -> tuple[bool, str]:
        """
        Export content as .txt file.
        
        Args:
            content (str): Content to export
            filepath (str): Where to save file
            
        Returns:
            tuple: (success: bool, message: str)
        """
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Exported .txt to: {filepath}")
            return True, f"Exported to {filepath}"
        except Exception as e:
            error = f"Export failed: {str(e)}"
            logger.error(error)
            return False, error
    
    def export_docx(self, content: str, filepath: str) -> tuple[bool, str]:
        """
        Export content as .docx file.
        
        Args:
            content (str): Content to export
            filepath (str): Where to save file
            
        Returns:
            tuple: (success: bool, message: str)
        """
        try:
            from docx import Document
            
            document = Document()
            document.add_heading('n8n Response Summary', 0)
            document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            document.add_paragraph()
            
            # Split content into paragraphs for better formatting
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    document.add_paragraph(paragraph)
            
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            document.save(filepath)
            
            logger.info(f"Exported .docx to: {filepath}")
            return True, f"Exported to {filepath}"
        except ImportError:
            error = "python-docx not installed. Install with: pip install python-docx"
            logger.error(error)
            return False, error
        except Exception as e:
            error = f"Export failed: {str(e)}"
            logger.error(error)
            return False, error
    
    def generate_filename(self, original_path: str, suffix: str = "_Summary", 
                         format: str = "txt") -> str:
        """
        Generate export filename intelligently.
        
        Args:
            original_path (str): Original file path
            suffix (str): Suffix to add before extension
            format (str): File format ('txt' or 'docx')
            
        Returns:
            str: New filename
            
        Example:
            >>> model = FileModel()
            >>> filename = model.generate_filename(
            ...     '/path/to/document.pdf',
            ...     suffix='_Summary',
            ...     format='txt'
            ... )
            >>> print(filename)
            '/path/to/document_Summary.txt'
        """
        base = os.path.splitext(original_path)[0]
        return f"{base}{suffix}.{format}"
    
    def generate_timestamp_filename(self, format: str = "txt") -> str:
        """
        Generate filename based on current timestamp.
        
        Args:
            format (str): File format ('txt' or 'docx')
            
        Returns:
            str: Filename like 'n8n_response_20250207_144235.txt'
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"n8n_response_{timestamp}.{format}"
